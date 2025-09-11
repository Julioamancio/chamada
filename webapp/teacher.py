from __future__ import annotations

from datetime import date, datetime, timedelta
from io import BytesIO, StringIO
import csv

from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    send_file,
    current_app,
    session,
)
from flask_login import login_required, current_user

from . import db
from .models import (
    Classroom,
    Student,
    Session,
    AttendanceEntry,
    get_or_create_session,
    Stage,
    StageSession,
    Activity,
    DailyScore,
)
from .calendar_utils import brazil_national_holidays, generate_teaching_dates

# Compat import: supports both new and old helper names
try:
    from .excel import names_from_upload as _get_names
except Exception:
    try:
        from .excel import names_from_first_column as _get_names
    except Exception:
        _get_names = None


bp = Blueprint("teacher", __name__, template_folder="templates")


def teacher_required(func):
    from functools import wraps

    @wraps(func)
    def wrapper(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for("auth.login"))
        if current_user.role not in ("teacher", "admin"):
            flash("Acesso restrito a professores.", "warning")
            return redirect(url_for("auth.login"))
        return func(*args, **kwargs)

    return wrapper


@bp.route("/")
@login_required
@teacher_required
def dashboard():
    if current_user.role == "admin":
        classes = Classroom.query.order_by(Classroom.name).all()
    else:
        classes = (
            Classroom.query.filter_by(owner_id=current_user.id)
            .order_by(Classroom.name)
            .all()
        )
    return render_template("teacher/classes.html", classes=classes)


@bp.route("/classes/create", methods=["POST"])
@login_required
@teacher_required
def create_class():
    name = request.form.get("name", "").strip()
    if not name:
        flash("Informe o nome da turma.", "warning")
        return redirect(url_for("teacher.dashboard"))
    c = Classroom(name=name, owner_id=current_user.id)
    db.session.add(c)
    db.session.commit()
    flash("Turma criada.", "success")
    return redirect(url_for("teacher.dashboard"))


@bp.route("/classes/<int:class_id>/delete", methods=["POST"])
@login_required
@teacher_required
def delete_class(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    db.session.delete(c)
    db.session.commit()
    flash("Turma excluida.", "success")
    return redirect(url_for("teacher.dashboard"))


@bp.route("/classes/<int:class_id>/rename", methods=["POST"])
@login_required
@teacher_required
def rename_class(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    new_name = request.form.get("name", "").strip()
    if not new_name:
        flash("Informe o novo nome.", "warning")
        return redirect(url_for("teacher.dashboard"))
    c.name = new_name
    db.session.commit()
    flash("Turma renomeada.", "success")
    return redirect(url_for("teacher.dashboard"))


@bp.route("/classes/<int:class_id>")
@login_required
@teacher_required
def class_detail(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    dt = request.args.get("date") or date.today().isoformat()
    students = (
        Student.query.filter_by(class_id=c.id).order_by(Student.name).all()
    )
    s = get_or_create_session(c, dt)
    # Garante que a sessão tente se vincular à etapa correspondente pela data
    try:
        _attach_stage_to_session(c, s, dt)
    except Exception:
        pass
    entries = {e.student_id: bool(e.present) for e in s.entries}
    stages = Stage.query.filter_by(class_id=c.id).order_by(Stage.name).all()
    return render_template(
        "teacher/class_detail.html",
        classroom=c,
        students=students,
        entries=entries,
        att_date=dt,
        stages=stages,
    )


@bp.route("/classes/<int:class_id>/students/add", methods=["POST"])
@login_required
@teacher_required
def add_student(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    name = request.form.get("name", "").strip()
    if not name:
        flash("Nome do aluno e obrigatorio.", "warning")
        return redirect(url_for("teacher.class_detail", class_id=class_id))
    st = Student(name=name, class_id=class_id)
    db.session.add(st)
    db.session.commit()
    flash("Aluno adicionado.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id))


@bp.route("/classes/<int:class_id>/students/<int:student_id>/delete", methods=["POST"])
@login_required
@teacher_required
def delete_student(class_id: int, student_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    st = Student.query.get_or_404(student_id)
    db.session.delete(st)
    db.session.commit()
    flash("Aluno removido.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id))


@bp.route("/classes/<int:class_id>/students/<int:student_id>/rename", methods=["POST"])
@login_required
@teacher_required
def rename_student(class_id: int, student_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    st = Student.query.get_or_404(student_id)
    new_name = request.form.get("name", "").strip()
    if not new_name:
        flash("Informe o novo nome do aluno.", "warning")
        return redirect(url_for("teacher.class_detail", class_id=class_id))
    st.name = new_name
    db.session.commit()
    flash("Aluno renomeado.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id))


@bp.route("/classes/<int:class_id>/attendance/save", methods=["POST"])
@login_required
@teacher_required
def save_attendance(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    # Admin ou dono da turma podem editar chamadas
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    dt = request.form.get("date") or date.today().isoformat()
    session = get_or_create_session(c, dt)
    _attach_stage_to_session(c, session, dt)
    present_ids = set(int(sid) for sid in request.form.getlist("present"))
    entries_map = {e.student_id: e for e in session.entries}
    for st in Student.query.filter_by(class_id=c.id).all():
        e = entries_map.get(st.id)
        val = st.id in present_ids
        if e:
            e.present = val
        else:
            db.session.add(
                AttendanceEntry(session_id=session.id, student_id=st.id, present=val)
            )
    db.session.commit()
    # Update activity scores for this date (idempotent upsert)
    try:
        _update_scores_for_date(class_id=c.id, dt=dt)
    except Exception:
        # never fail saving attendance due to scoring update
        pass
    # For fetch/AJAX autosave, avoid redirect payload
    if request.headers.get("X-Requested-With"):
        return ("", 204)
    flash("Chamada salva.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id, date=dt))


@bp.route("/classes/<int:class_id>/calendar", methods=["GET", "POST"])
@login_required
@teacher_required
def class_calendar(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))

    stages = Stage.query.filter_by(class_id=c.id).order_by(Stage.name).all()
    try:
        from .models import GlobalStage  # local import if not already
        global_stages = GlobalStage.query.order_by(GlobalStage.name).all()
    except Exception:
        global_stages = []
    ctx = {"classroom": c, "stages": stages, "global_stages": global_stages}
    return render_template("teacher/calendar.html", **ctx)


@bp.route("/classes/<int:class_id>/stages/create", methods=["POST"])
@login_required
@teacher_required
def create_stage(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    name = (request.form.get("stage_name") or "").strip()
    start = request.form.get("start")
    end = request.form.get("end")
    weekdays = ",".join(sorted(set(request.form.getlist("weekday"))))
    if not name or not start or not end or not weekdays:
        flash("Preencha nome, periodo e dia(s) da semana.", "warning")
        return redirect(url_for("teacher.class_calendar", class_id=class_id))
    stg = Stage(class_id=c.id, name=name, start=start, end=end, weekdays=weekdays)
    db.session.add(stg)
    db.session.commit()
    flash("Etapa criada.", "success")
    return redirect(url_for("teacher.class_calendar", class_id=class_id))


@bp.route("/classes/<int:class_id>/stages/<int:stage_id>/delete", methods=["POST"])
@login_required
@teacher_required
def delete_stage(class_id: int, stage_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    stg = Stage.query.get_or_404(stage_id)
    db.session.delete(stg)
    db.session.commit()
    flash("Etapa excluida.", "success")
    return redirect(url_for("teacher.class_calendar", class_id=class_id))


@bp.route("/classes/<int:class_id>/stages/<int:stage_id>/edit", methods=["POST"])
@login_required
@teacher_required
def update_stage(class_id: int, stage_id: int):
    c = Classroom.query.get_or_404(class_id)
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    stg = Stage.query.get_or_404(stage_id)
    name = (request.form.get("stage_name") or "").strip()
    start = request.form.get("start")
    end = request.form.get("end")
    weekdays = ",".join(sorted(set(request.form.getlist("weekday"))))
    if not name or not start or not end or not weekdays:
        flash("Preencha nome, periodo e dia(s) da semana.", "warning")
        return redirect(url_for("teacher.class_calendar", class_id=class_id))
    stg.name = name
    stg.start = start
    stg.end = end
    stg.weekdays = weekdays
    db.session.commit()
    flash("Etapa atualizada.", "success")
    return redirect(url_for("teacher.class_calendar", class_id=class_id))


@bp.route("/classes/<int:class_id>/calls")
@login_required
@teacher_required
def calls(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    stage_id = request.args.get("stage_id", type=int)
    q = Session.query.filter_by(class_id=c.id)
    start = end = None
    if stage_id:
        stg = Stage.query.get_or_404(stage_id)
        start, end = stg.start, stg.end
        q = q.filter(Session.date >= start, Session.date <= end)
    sessions = q.order_by(Session.date.desc()).all()
    # Build counts per session + stage name/id
    rows = []
    total_students = Student.query.filter_by(class_id=c.id).count()
    # Map session -> stage via StageSession when available
    sess_ids = [s.id for s in sessions]
    mapping = {}
    if sess_ids:
        try:
            mapping = {m.session_id: m.stage_id for m in StageSession.query.filter(StageSession.session_id.in_(sess_ids)).all()}
        except Exception:
            mapping = {}
    stage_by_id = {st.id: st for st in Stage.query.filter_by(class_id=c.id).all()}
    for s in sessions:
        pres = sum(1 for e in s.entries if e.present)
        # Resolve stage name: from mapping or by date fallback
        st_name = None
        st_id = mapping.get(s.id)
        if st_id and stage_by_id.get(st_id):
            st_name = stage_by_id[st_id].name
        else:
            st = _find_stage_for_date(c, s.date)
            if st:
                st_name = st.name
                st_id = st.id
        rows.append({
            "session_id": s.id,
            "date": s.date,
            "present": pres,
            "absent": total_students - pres,
            "stage": st_name,
            "stage_id": st_id,
        })
    stages = Stage.query.filter_by(class_id=c.id).order_by(Stage.name).all()
    return render_template("teacher/calls.html", classroom=c, stages=stages, rows=rows, selected_stage=stage_id)


def _find_stage_for_date(classroom: Classroom, dt_str: str):
    try:
        d = datetime.strptime(dt_str, "%Y-%m-%d").date()
    except Exception:
        return None
    stgs = Stage.query.filter_by(class_id=classroom.id).all()
    for stg in stgs:
        try:
            s = datetime.strptime(stg.start, "%Y-%m-%d").date()
            e = datetime.strptime(stg.end, "%Y-%m-%d").date()
        except Exception:
            continue
        if s <= d <= e:
            if stg.weekdays:
                try:
                    w = set(int(x) for x in stg.weekdays.split(",") if x)
                except Exception:
                    w = set()
                if w and d.weekday() not in w:
                    continue
            return stg
    return None


def _attach_stage_to_session(classroom: Classroom, sess: Session, dt_str: str, force_stage: Stage | None = None):
    try:
        from .models import StageSession  # local import in case of circular
        current = StageSession.query.filter_by(session_id=sess.id).first()
        target = force_stage or _find_stage_for_date(classroom, dt_str)
        if not target:
            return
        if not current:
            db.session.add(StageSession(session_id=sess.id, stage_id=target.id))
        elif current.stage_id != target.id:
            # Corrige mapeamento se a data indicar outra etapa
            current.stage_id = target.id
    except Exception:
        return


@bp.route("/classes/<int:class_id>/calls/<int:session_id>/set_stage", methods=["POST"])
@login_required
@teacher_required
def set_session_stage(class_id: int, session_id: int):
    """Define manualmente a etapa de uma chamada (sessão).
    Cria ou atualiza o vínculo em StageSession. Se "stage_id" vier vazio,
    tenta remover o vínculo existente.
    """
    c = Classroom.query.get_or_404(class_id)
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    sess = Session.query.get_or_404(session_id)
    try:
        from .models import StageSession
        stage_id = request.form.get("stage_id", type=int)
        mapping = StageSession.query.filter_by(session_id=sess.id).first()
        if stage_id:
            # valida se a etapa pertence à turma
            stg = Stage.query.get_or_404(stage_id)
            if stg.class_id != c.id:
                flash("Etapa inválida para esta turma.", "warning")
                return redirect(url_for("teacher.calls", class_id=class_id))
            if mapping:
                mapping.stage_id = stage_id
            else:
                db.session.add(StageSession(session_id=sess.id, stage_id=stage_id))
        else:
            # remover vínculo
            if mapping:
                db.session.delete(mapping)
        db.session.commit()
        flash("Etapa da chamada atualizada.", "success")
    except Exception:
        flash("Falha ao atualizar etapa da chamada.", "danger")
    return redirect(url_for("teacher.calls", class_id=class_id))


@bp.route("/classes/<int:class_id>/calls/<int:session_id>/delete", methods=["POST"])
@login_required
@teacher_required
def delete_call(class_id: int, session_id: int):
    """Exclui a chamada (sessão) e seus vínculos/entradas."""
    c = Classroom.query.get_or_404(class_id)
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    sess = Session.query.get_or_404(session_id)
    if sess.class_id != c.id:
        flash("Chamada inválida.", "warning")
        return redirect(url_for("teacher.calls", class_id=class_id))
    try:
        from .models import StageSession
        # remover mapeamento etapa -> sessão se existir
        m = StageSession.query.filter_by(session_id=sess.id).first()
        if m:
            db.session.delete(m)
        # remover a sessão (AttendanceEntry tem cascade)
        db.session.delete(sess)
        db.session.commit()
        flash("Chamada excluída.", "success")
    except Exception:
        flash("Falha ao excluir chamada.", "danger")
    return redirect(url_for("teacher.calls", class_id=class_id))


@bp.route("/set-lang", methods=["POST"])
@login_required
@teacher_required
def set_lang():
    """Atualiza o idioma preferido do usuário na sessão.
    Suporta: 'pt-br' (Português), 'en' (English), 'zh' (中文).
    """
    lang = (request.form.get("lang") or "pt-br").lower()
    if lang not in ("pt-br", "en", "zh"):
        lang = "pt-br"
    session["lang"] = lang
    # evita mensagem para não poluir a UI; apenas redireciona
    return redirect(request.referrer or url_for("teacher.dashboard"))


@bp.route("/classes/<int:class_id>/stages/link", methods=["POST"])
@login_required
@teacher_required
def link_global_stages(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    from .models import GlobalStage
    ids = [int(x) for x in request.form.getlist("template_id")]
    if not ids:
        flash("Selecione ao menos uma etapa global.", "warning")
        return redirect(url_for("teacher.class_calendar", class_id=class_id))
    created = 0
    for tid in ids:
        tpl = GlobalStage.query.get_or_404(tid)
        exists = Stage.query.filter_by(class_id=c.id, name=tpl.name, start=tpl.start, end=tpl.end).first()
        if exists:
            continue
        stg = Stage(class_id=c.id, name=tpl.name, start=tpl.start, end=tpl.end, weekdays=tpl.weekdays)
        db.session.add(stg)
        created += 1
    db.session.commit()
    flash(f"{created} etapa(s) vinculadas à turma.", "success")
    return redirect(url_for("teacher.class_calendar", class_id=class_id))


# Global stages (accessible to professores e administradores)
@bp.route("/stages")
@login_required
@teacher_required
def global_stages():
    from .models import GlobalStage
    stages = GlobalStage.query.order_by(GlobalStage.name).all()
    return render_template("teacher/stages.html", stages=stages)


@bp.route("/stages/create", methods=["POST"])
@login_required
@teacher_required
def create_global_stage():
    from .models import GlobalStage
    name = (request.form.get("stage_name") or "").strip()
    start = request.form.get("start")
    end = request.form.get("end")
    weekdays = ",".join(sorted(set(request.form.getlist("weekday"))))
    if not name or not start or not end or not weekdays:
        flash("Preencha nome, periodo e dia(s) da semana.", "warning")
        return redirect(url_for("teacher.global_stages"))
    stg = GlobalStage(name=name, start=start, end=end, weekdays=weekdays)
    db.session.add(stg)
    db.session.commit()
    flash("Etapa global criada.", "success")
    return redirect(url_for("teacher.global_stages"))


@bp.route("/stages/<int:stage_id>/edit", methods=["POST"])
@login_required
@teacher_required
def edit_global_stage(stage_id: int):
    from .models import GlobalStage
    stg = GlobalStage.query.get_or_404(stage_id)
    name = (request.form.get("stage_name") or "").strip()
    start = request.form.get("start")
    end = request.form.get("end")
    weekdays = ",".join(sorted(set(request.form.getlist("weekday"))))
    if not name or not start or not end or not weekdays:
        flash("Preencha nome, periodo e dia(s) da semana.", "warning")
        return redirect(url_for("teacher.global_stages"))
    stg.name, stg.start, stg.end, stg.weekdays = name, start, end, weekdays
    db.session.commit()
    flash("Etapa global atualizada.", "success")
    return redirect(url_for("teacher.global_stages"))


@bp.route("/stages/<int:stage_id>/delete", methods=["POST"])
@login_required
@teacher_required
def delete_global_stage(stage_id: int):
    from .models import GlobalStage
    stg = GlobalStage.query.get_or_404(stage_id)
    db.session.delete(stg)
    db.session.commit()
    flash("Etapa global excluida.", "success")
    return redirect(url_for("teacher.global_stages"))


@bp.route("/classes/<int:class_id>/attendance/bulk", methods=["POST"])
@login_required
@teacher_required
def bulk_attendance(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    # Admin ou dono da turma podem editar chamadas
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    stage_id = request.form.get("stage_id", type=int)
    stg = Stage.query.get_or_404(stage_id)
    present_ids = set(int(sid) for sid in request.form.getlist("present"))
    start = datetime.strptime(stg.start, "%Y-%m-%d").date()
    end = datetime.strptime(stg.end, "%Y-%m-%d").date()
    wd = [int(x) for x in stg.weekdays.split(",") if x]
    holidays = set()
    for y in {start.year, end.year}:
        holidays |= brazil_national_holidays(y)
    dates = generate_teaching_dates(start, end, wd, holidays)
    for d in dates:
        sess = get_or_create_session(c, d.isoformat())
        _attach_stage_to_session(c, sess, d.isoformat(), force_stage=stg)
        entries_map = {e.student_id: e for e in sess.entries}
        for st in Student.query.filter_by(class_id=c.id).all():
            val = st.id in present_ids
            e = entries_map.get(st.id)
            if e:
                e.present = val
            else:
                db.session.add(
                    AttendanceEntry(session_id=sess.id, student_id=st.id, present=val)
                )
    db.session.commit()
    flash("Chamada em lote aplicada a etapa.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id))


@bp.route("/classes/<int:class_id>/attendance/bulk-range", methods=["POST"])
@login_required
@teacher_required
def bulk_attendance_range(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    # Admin ou dono da turma podem editar chamadas
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    start_s = request.form.get("start")
    end_s = request.form.get("end")
    if not start_s or not end_s:
        flash("Informe periodo inicio e fim.", "warning")
        return redirect(url_for("teacher.class_detail", class_id=class_id))
    start = datetime.strptime(start_s, "%Y-%m-%d").date()
    end = datetime.strptime(end_s, "%Y-%m-%d").date()
    if end < start:
        start, end = end, start
    present_ids = set(int(sid) for sid in request.form.getlist("present"))
    d = start
    while d <= end:
        sess = get_or_create_session(c, d.isoformat())
        _attach_stage_to_session(c, sess, d.isoformat())
        entries_map = {e.student_id: e for e in sess.entries}
        for st in Student.query.filter_by(class_id=c.id).all():
            val = st.id in present_ids
            e = entries_map.get(st.id)
            if e:
                e.present = val
            else:
                db.session.add(
                    AttendanceEntry(session_id=sess.id, student_id=st.id, present=val)
                )
        d += timedelta(days=1)
    db.session.commit()
    flash("Chamada em lote aplicada ao periodo.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id))


@bp.route("/classes/<int:class_id>/attendance/auto-range", methods=["POST"])
@login_required
@teacher_required
def bulk_attendance_auto_range(class_id: int):
    """Apply randomized attendance over a date range at a target rate.

    - rate: fraction of students marked present per day (default 0.8)
    - Ensures no student is absent on two consecutive days (best-effort)
    """
    import random

    c = Classroom.query.get_or_404(class_id)
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))

    start_s = request.form.get("start")
    end_s = request.form.get("end")
    rate = request.form.get("rate", type=float) or 0.8
    if not start_s or not end_s:
        flash("Informe período início e fim.", "warning")
        return redirect(url_for("teacher.class_detail", class_id=class_id))

    start = datetime.strptime(start_s, "%Y-%m-%d").date()
    end = datetime.strptime(end_s, "%Y-%m-%d").date()
    if end < start:
        start, end = end, start

    students = Student.query.filter_by(class_id=c.id).order_by(Student.name).all()
    ids = [s.id for s in students]
    n = len(ids)
    if n == 0:
        flash("Sem alunos para aplicar.", "warning")
        return redirect(url_for("teacher.class_detail", class_id=class_id))
    # number absent per day
    daily_abs = max(0, int(round(n * (1.0 - rate))))

    # Rotate absences so no student is absent in consecutive days
    last_absent: set[int] = set()

    d = start
    while d <= end:
        pool = [x for x in ids if x not in last_absent]
        if len(pool) < daily_abs:
            # If not enough, allow previous absences back but shuffle
            pool = ids[:]
        random.shuffle(pool)
        today_absent = set(pool[:daily_abs])
        today_present = set(ids) - today_absent

        sess = get_or_create_session(c, d.isoformat())
        _attach_stage_to_session(c, sess, d.isoformat())
        entries_map = {e.student_id: e for e in sess.entries}
        for sid in ids:
            val = sid in today_present
            e = entries_map.get(sid)
            if e:
                e.present = val
            else:
                db.session.add(AttendanceEntry(session_id=sess.id, student_id=sid, present=val))
        last_absent = today_absent
        d += timedelta(days=1)

    db.session.commit()
    flash("Chamada aleatória (80%) aplicada ao período.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id))


@bp.route("/classes/<int:class_id>/attendance/auto-stage", methods=["POST"])
@login_required
@teacher_required
def bulk_attendance_auto_stage(class_id: int):
    import random

    c = Classroom.query.get_or_404(class_id)
    if not (current_user.role == "admin" or c.owner_id == current_user.id):
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))

    stage_id = request.form.get("stage_id", type=int)
    rate = request.form.get("rate", type=float) or 0.8
    stg = Stage.query.get_or_404(stage_id)

    start = datetime.strptime(stg.start, "%Y-%m-%d").date()
    end = datetime.strptime(stg.end, "%Y-%m-%d").date()
    wd = [int(x) for x in stg.weekdays.split(",") if x]
    holidays = set()
    for y in {start.year, end.year}:
        holidays |= brazil_national_holidays(y)
    dates = generate_teaching_dates(start, end, wd, holidays)

    students = Student.query.filter_by(class_id=c.id).order_by(Student.name).all()
    ids = [s.id for s in students]
    n = len(ids)
    daily_abs = max(0, int(round(n * (1.0 - rate))))
    last_absent: set[int] = set()

    for d in dates:
        pool = [x for x in ids if x not in last_absent]
        if len(pool) < daily_abs:
            pool = ids[:]
        random.shuffle(pool)
        today_absent = set(pool[:daily_abs])
        today_present = set(ids) - today_absent
        sess = get_or_create_session(c, d.isoformat())
        _attach_stage_to_session(c, sess, d.isoformat(), force_stage=stg)
        entries_map = {e.student_id: e for e in sess.entries}
        for sid in ids:
            val = sid in today_present
            e = entries_map.get(sid)
            if e:
                e.present = val
            else:
                db.session.add(AttendanceEntry(session_id=sess.id, student_id=sid, present=val))
        last_absent = today_absent
    db.session.commit()
    flash("Chamada aleatoria (80%) aplicada na etapa.", "success")
    return redirect(url_for("teacher.class_detail", class_id=class_id))

@bp.route("/classes/<int:class_id>/reports")
@login_required
@teacher_required
def reports(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissão.", "danger")
        return redirect(url_for("teacher.dashboard"))
    students = Student.query.filter_by(class_id=c.id).order_by(Student.name).all()
    stages = Stage.query.filter_by(class_id=c.id).order_by(Stage.name).all()
    return render_template("teacher/reports.html", classroom=c, students=students, stages=stages)


# =====================
# Settings for teacher (theme + backups)

@bp.route("/settings")
@login_required
@teacher_required
def settings():
    """Teacher settings page with theme toggle and personal backup tools."""
    return render_template("teacher/settings.html")


@bp.route("/settings/export", methods=["GET"])
@login_required
@teacher_required
def export_teacher_backup():
    from io import BytesIO
    import json
    from datetime import datetime

    classes = Classroom.query.filter_by(owner_id=current_user.id).order_by(Classroom.name).all()
    payload = {
        "version": 1,
        "teacher": {"id": current_user.id, "name": current_user.name, "email": current_user.email},
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "classes": [],
    }
    for c in classes:
        # stages
        stages = [
            {"name": s.name, "start": s.start, "end": s.end, "weekdays": s.weekdays}
            for s in Stage.query.filter_by(class_id=c.id).order_by(Stage.name).all()
        ]
        # students
        students = [
            {"name": s.name}
            for s in Student.query.filter_by(class_id=c.id).order_by(Student.name).all()
        ]
        # sessions + entries
        sessions = []
        # map session->stage via StageSession
        try:
            mapping = {m.session_id: m.stage_id for m in StageSession.query.join(Session, StageSession.session_id == Session.id).filter(Session.class_id == c.id).all()}
        except Exception:
            mapping = {}
        stage_by_id = {s.id: s for s in Stage.query.filter_by(class_id=c.id).all()}
        for sess in Session.query.filter_by(class_id=c.id).order_by(Session.date).all():
            entries = [
                {"student_name": e.student.name, "present": bool(e.present)}
                for e in sorted(sess.entries, key=lambda x: x.student.name)
            ]
            stg_name = None
            sid = mapping.get(sess.id)
            if sid and stage_by_id.get(sid):
                stg_name = stage_by_id[sid].name
            sessions.append({"date": sess.date, "stage_name": stg_name, "entries": entries})
        payload["classes"].append({
            "name": c.name,
            "stages": stages,
            "students": students,
            "sessions": sessions,
        })

    data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    bio = BytesIO(data); bio.seek(0)
    fname = f"backup_prof_{current_user.name.replace(' ','_')}.json"
    return send_file(bio, as_attachment=True, download_name=fname, mimetype="application/json; charset=utf-8")


@bp.route("/settings/import", methods=["POST"])
@login_required
@teacher_required
def import_teacher_backup():
    import json
    from datetime import datetime

    f = request.files.get("file")
    if not f or not f.filename.lower().endswith(".json"):
        flash("Envie um arquivo .json de backup válido.", "warning")
        return redirect(url_for("teacher.settings"))
    try:
        payload = json.loads(f.read().decode("utf-8"))
    except Exception as e:
        flash(f"Falha ao ler JSON: {e}", "danger")
        return redirect(url_for("teacher.settings"))

    classes = payload.get("classes", [])
    imported = 0
    ts = datetime.now().strftime("%Y%m%d-%H%M")
    for cdata in classes:
        base_name = cdata.get("name") or "Turma"
        name = base_name
        # avoid name conflict
        if Classroom.query.filter_by(owner_id=current_user.id, name=name).first():
            name = f"{base_name} (import {ts})"
        c = Classroom(name=name, owner_id=current_user.id)
        db.session.add(c)
        db.session.flush()
        # stages
        stage_map = {}
        for s in cdata.get("stages", []):
            stg = Stage(class_id=c.id, name=s.get("name",""), start=s.get("start",""), end=s.get("end",""), weekdays=s.get("weekdays",""))
            db.session.add(stg); db.session.flush(); stage_map[stg.name] = stg.id
        # students
        student_map = {}
        for s in cdata.get("students", []):
            st = Student(name=s.get("name",""), class_id=c.id)
            db.session.add(st); db.session.flush(); student_map[st.name] = st.id
        # sessions + entries
        for s in cdata.get("sessions", []):
            sess = get_or_create_session(c, s.get("date"))
            entries_map = {e.student_id: e for e in sess.entries}
            for e in s.get("entries", []):
                sid = student_map.get(e.get("student_name"))
                if not sid: continue
                present = bool(e.get("present"))
                ent = entries_map.get(sid)
                if ent:
                    ent.present = present
                else:
                    db.session.add(AttendanceEntry(session_id=sess.id, student_id=sid, present=present))
            # stage mapping by name
            stg_name = s.get("stage_name")
            if stg_name and stage_map.get(stg_name):
                try:
                    if not StageSession.query.filter_by(session_id=sess.id).first():
                        db.session.add(StageSession(session_id=sess.id, stage_id=stage_map[stg_name]))
                except Exception:
                    pass
        imported += 1
    db.session.commit()
    flash(f"Backup importado: {imported} turma(s) adicionadas.", "success")
    return redirect(url_for("teacher.settings"))


# =====================
# Activities (Atividades)

def _count_sessions_in_period(class_id: int, start: str, end: str, stage_id: int | None = None) -> int:
    """Return planned N (aulas) for the interval using stage weekdays.

    If stage has weekdays set, compute count by calendar (exclui feriados).
    Otherwise, fallback to counting sessions in DB.
    """
    try:
        stg = Stage.query.get(stage_id) if stage_id else None
        if stg and stg.weekdays:
            from .calendar_utils import brazil_national_holidays, generate_teaching_dates
            from datetime import datetime as _dt
            s = _dt.strptime(start, "%Y-%m-%d").date()
            e = _dt.strptime(end, "%Y-%m-%d").date()
            wd = [int(x) for x in stg.weekdays.split(",") if x]
            holidays = set()
            for y in {s.year, e.year}:
                holidays |= brazil_national_holidays(y)
            return len(generate_teaching_dates(s, e, wd, holidays))
    except Exception:
        pass
    q = Session.query.filter_by(class_id=class_id).filter(Session.date >= start, Session.date <= end)
    return q.count()


def _update_scores_for_activity(activity: Activity, class_id: int) -> None:
    """Populate/refresh DailyScore for an activity based on attendance entries.

    Idempotent: uses unique (activity_id, student_id, date) to upsert.
    """
    start, end = activity.period_start, activity.period_end
    q = Session.query.filter_by(class_id=class_id).filter(Session.date >= start, Session.date <= end)
    cand = q.order_by(Session.date).all()
    # Prefer sessions vinculadas à etapa. Complementa com datas que caem na etapa por data/semana.
    sess_ids = [s.id for s in cand]
    sessions = cand
    if sess_ids:
        mapped_ids = {
            m.session_id
            for m in StageSession.query.filter(
                StageSession.session_id.in_(sess_ids), StageSession.stage_id == activity.stage_id
            ).all()
        }
        if mapped_ids:
            by_map = {s.id: s for s in cand if s.id in mapped_ids}
            # Add sessions that fall into the activity's stage by date/weekdays
            extra = []
            try:
                stg = Stage.query.get(activity.stage_id)
                for s in cand:
                    if s.id in by_map:
                        continue
                    stg_by_date = _find_stage_for_date(s.classroom, s.date) if hasattr(s, 'classroom') else _find_stage_for_date(Classroom.query.get(class_id), s.date)
                    if stg_by_date and stg_by_date.id == activity.stage_id:
                        extra.append(s)
            except Exception:
                extra = []
            sessions = list(by_map.values()) + extra
    students = Student.query.filter_by(class_id=class_id).all()
    ppc = float(activity.points_per_call)
    for sess in sessions:
        entries = {e.student_id: bool(e.present) for e in sess.entries}
        for st in students:
            present = bool(entries.get(st.id))
            pts = ppc if present else 0.0
            row = DailyScore.query.filter_by(activity_id=activity.id, student_id=st.id, date=sess.date).first()
            if row:
                row.present = present
                row.points = pts
            else:
                db.session.add(DailyScore(activity_id=activity.id, student_id=st.id, class_id=class_id, date=sess.date, present=present, points=pts))
    db.session.commit()


def _update_scores_for_date(class_id: int, dt: str) -> None:
    """Update DailyScore rows for all activities overlapping date dt.
    Called when attendance is saved for a date.
    """
    # Find activities whose stage belongs to this class and whose period includes dt and are ATIVA
    acts = (
        db.session.query(Activity)
        .join(Stage, Activity.stage_id == Stage.id)
        .filter(
            Stage.class_id == class_id,
            Activity.status == "ATIVA",
            Activity.period_start <= dt,
            Activity.period_end >= dt,
        )
        .all()
    )
    if not acts:
        return
    # Load session and entries for the target date
    sess = Session.query.filter_by(class_id=class_id, date=dt).first()
    if not sess:
        return  # nothing to update without a session
    entries = {e.student_id: bool(e.present) for e in sess.entries}
    # If the session is explicitly mapped to a stage, honor that mapping
    try:
        mapped = None
        from .models import StageSession  # local import to avoid circular
        mapped = StageSession.query.filter_by(session_id=sess.id).first()
    except Exception:
        mapped = None
    students = Student.query.filter_by(class_id=class_id).all()
    for act in acts:
        # If this session is mapped to a different stage, skip this activity
        if mapped and mapped.stage_id != act.stage_id:
            continue
        ppc = float(act.points_per_call)
        for st in students:
            present = bool(entries.get(st.id))
            pts = ppc if present else 0.0
            row = (
                DailyScore.query.filter_by(activity_id=act.id, student_id=st.id, date=dt)
                .first()
            )
            if row:
                row.present = present
                row.points = pts
            else:
                db.session.add(
                    DailyScore(
                        activity_id=act.id,
                        student_id=st.id,
                        class_id=class_id,
                        date=dt,
                        present=present,
                        points=pts,
                    )
                )
    db.session.commit()


@bp.route("/classes/<int:class_id>/activities")
@login_required
@teacher_required
def activities(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissão.", "danger")
        return redirect(url_for("teacher.dashboard"))
    stages = Stage.query.filter_by(class_id=c.id).order_by(Stage.name).all()
    stage_id = request.args.get("stage_id", type=int)
    q = (
        Activity.query.join(Stage, Activity.stage_id == Stage.id)
        .filter(Stage.class_id == c.id)
        .order_by(Activity.created_at.desc())
    )
    if stage_id:
        q = q.filter(Activity.stage_id == stage_id)
    acts = q.all()
    # annotate with real_n
    annotated = []
    for a in acts:
        real_n = _count_sessions_in_period(class_id=c.id, start=a.period_start, end=a.period_end, stage_id=a.stage_id)
        annotated.append((a, real_n))
    # preview data per stage for the creation form
    preview = {}
    for stg in stages:
        n = _count_sessions_in_period(class_id=c.id, start=stg.start, end=stg.end, stage_id=stg.id)
        preview[stg.id] = {"n": n, "start": stg.start, "end": stg.end}
    return render_template(
        "teacher/activities.html",
        classroom=c,
        stages=stages,
        activities=annotated,
        selected_stage=stage_id,
        preview_by_stage=preview,
    )


@bp.route("/classes/<int:class_id>/activities/create", methods=["POST"])
@login_required
@teacher_required
def create_activity(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissão.", "danger")
        return redirect(url_for("teacher.dashboard"))
    stage_id = request.form.get("stage_id", type=int)
    title = (request.form.get("title") or "").strip()
    desc = (request.form.get("description") or "").strip() or None
    total = request.form.get("points_total", type=float)
    if not (stage_id and title and total):
        flash("Preencha etapa, título e valor total.", "warning")
        return redirect(url_for("teacher.activities", class_id=class_id))
    stg = Stage.query.get_or_404(stage_id)
    start, end = stg.start, stg.end
    lessons = _count_sessions_in_period(class_id=c.id, start=start, end=end, stage_id=stage_id)
    if not lessons or lessons <= 0:
        flash("Quantidade de aulas inválida.", "warning")
        return redirect(url_for("teacher.activities", class_id=class_id))
    ppc = round(total / lessons, 4)
    act = Activity(
        stage_id=stage_id,
        title=title,
        description=desc,
        period_start=start,
        period_end=end,
        lessons_count=lessons,
        points_total=total,
        points_per_call=ppc,
        created_by_user_id=current_user.id,
        created_by_role=("ADM" if current_user.role == "admin" else "PROFESSOR"),
        status="ATIVA",
    )
    db.session.add(act)
    db.session.commit()
    _update_scores_for_activity(act, class_id=c.id)
    flash("Atividade criada e pontuações calculadas.", "success")
    return redirect(url_for("teacher.activities", class_id=class_id))


@bp.route("/activities/<int:activity_id>/edit", methods=["POST"])
@login_required
@teacher_required
def edit_activity(activity_id: int):
    act = Activity.query.get_or_404(activity_id)
    stg = Stage.query.get_or_404(act.stage_id)
    c = Classroom.query.get_or_404(stg.class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissão.", "danger")
        return redirect(url_for("teacher.dashboard"))
    title = (request.form.get("title") or act.title).strip()
    desc = (request.form.get("description") or "").strip() or None
    total = request.form.get("points_total", type=float) or float(act.points_total)
    status = request.form.get("status") or act.status
    # Period and lessons are determined by the stage
    stg = Stage.query.get_or_404(act.stage_id)
    start, end = stg.start, stg.end
    lessons = _count_sessions_in_period(class_id=c.id, start=start, end=end, stage_id=act.stage_id)
    if lessons <= 0:
        flash("Quantidade de aulas inválida.", "warning")
        return redirect(url_for("teacher.activities", class_id=c.id))
    ppc = round(total / lessons, 4)
    act.title = title
    act.description = desc
    act.period_start = start
    act.period_end = end
    act.points_total = total
    act.lessons_count = lessons
    act.points_per_call = ppc
    act.status = status
    db.session.commit()
    # Harmoniza mapeamentos de sessões do período com a etapa da atividade
    try:
        q = Session.query.filter_by(class_id=c.id).filter(Session.date >= start, Session.date <= end)
        for sess in q.all():
            _attach_stage_to_session(c, sess, sess.date, force_stage=stg)
        db.session.commit()
    except Exception:
        pass
    _update_scores_for_activity(act, class_id=c.id)
    flash("Atividade atualizada e pontuações recalculadas.", "success")
    return redirect(url_for("teacher.activities", class_id=c.id))


@bp.route("/activities/<int:activity_id>/recalc", methods=["POST"])
@login_required
@teacher_required
def recalc_activity(activity_id: int):
    act = Activity.query.get_or_404(activity_id)
    stg = Stage.query.get_or_404(act.stage_id)
    c = Classroom.query.get_or_404(stg.class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissão.", "danger")
        return redirect(url_for("teacher.dashboard"))
    # Ajusta mapeamento das sessões do período para a etapa da atividade
    try:
        q = Session.query.filter_by(class_id=c.id).filter(Session.date >= stg.start, Session.date <= stg.end)
        for sess in q.all():
            _attach_stage_to_session(c, sess, sess.date, force_stage=stg)
        db.session.commit()
    except Exception:
        pass
    _update_scores_for_activity(act, class_id=c.id)
    flash("Pontuações recalculadas a partir das chamadas.", "success")
    return redirect(url_for("teacher.activities", class_id=c.id))


@bp.route("/classes/<int:class_id>/import", methods=["GET", "POST"])
@login_required
@teacher_required
def import_students(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))
    if request.method == "POST":
        if _get_names is None:
            flash("Falha do importador. Reinicie o servidor e tente novamente.", "danger")
            return redirect(url_for("teacher.import_students", class_id=class_id))
        file = request.files.get("file")
        if not file or not file.filename:
            flash("Envie um arquivo .xlsx ou .csv com nomes na coluna A.", "warning")
            return redirect(url_for("teacher.import_students", class_id=class_id))
        try:
            names = _get_names(file)
        except Exception as e:
            flash(f"Erro ao ler arquivo: {e}", "danger")
            return redirect(url_for("teacher.import_students", class_id=class_id))
        count = 0
        for n in names:
            n = n.strip()
            if not n:
                continue
            if not Student.query.filter_by(class_id=class_id, name=n).first():
                db.session.add(Student(name=n, class_id=class_id))
                count += 1
        db.session.commit()
        flash(f"Importacao concluida: {count} aluno(s) adicionados.", "success")
        return redirect(url_for("teacher.class_detail", class_id=class_id))
    return render_template("teacher/import.html", classroom=c)


@bp.route("/classes/<int:class_id>/export.csv")
@login_required
@teacher_required
def export_csv(class_id: int):
    c = Classroom.query.get_or_404(class_id)
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))

    # Write CSV to text stream then encode to bytes (Windows-safe)
    text = StringIO()
    writer = csv.writer(text)
    writer.writerow(["date", "student", "present"])
    sessions = Session.query.filter_by(class_id=c.id).order_by(Session.date).all()
    students = {s.id: s.name for s in Student.query.filter_by(class_id=c.id).all()}
    for s in sessions:
        entries = {e.student_id: e.present for e in s.entries}
        for sid, name in students.items():
            writer.writerow([s.date, name, "yes" if entries.get(sid) else "no"])
    data = text.getvalue().encode("utf-8")
    bio = BytesIO(data)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"{c.name}.csv",
        mimetype="text/csv; charset=utf-8",
    )


# Reports: DOCX exports (class and individual)
def _brand_logo_path():
    from pathlib import Path
    folder = Path(current_app.instance_path) / "branding"
    for p in folder.glob("logo.*"):
        return str(p)
    return None


def _attendance_summary_for_class(c: Classroom, start: str | None = None, end: str | None = None):
    # Build per-student totals
    students = Student.query.filter_by(class_id=c.id).order_by(Student.name).all()
    q = Session.query.filter_by(class_id=c.id)
    if start:
        q = q.filter(Session.date >= start)
    if end:
        q = q.filter(Session.date <= end)
    sessions = q.order_by(Session.date).all()
    totals = {s.id: {"name": s.name, "present": 0, "absent": 0} for s in students}
    for sess in sessions:
        entries = {e.student_id: bool(e.present) for e in sess.entries}
        for s in students:
            if entries.get(s.id):
                totals[s.id]["present"] += 1
            else:
                totals[s.id]["absent"] += 1
    return students, sessions, totals


@bp.route("/classes/<int:class_id>/report.docx")
@login_required
@teacher_required
def export_class_docx(class_id: int):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception:
        flash("Dependencia ausente: instale 'python-docx'.", "danger")
        return redirect(url_for("teacher.class_detail", class_id=class_id))

    c = Classroom.query.get_or_404(class_id)
    # View permission: admin can export, teacher owner can export
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))

    # Optional stage filter
    stage_id = request.args.get("stage_id", type=int)
    start = end = None
    if stage_id:
        stg = Stage.query.get_or_404(stage_id)
        start, end = stg.start, stg.end

    students, sessions, totals = _attendance_summary_for_class(c, start, end)

    doc = Document()
    # Header with logo and title
    section = doc.sections[0]
    section.left_margin, section.right_margin = Inches(0.7), Inches(0.7)
    header = section.header
    h_p = header.paragraphs[0]
    logo = _brand_logo_path()
    if logo:
        run = h_p.add_run()
        try:
            run.add_picture(logo, width=Inches(1.0))
        except Exception:
            pass

    title = doc.add_paragraph()
    r = title.add_run(f"Relatorio de Frequencia — {c.name}")
    r.font.size = Pt(16)
    r.bold = True

    period_txt = f" | Período: {start} → {end}" if start and end else ""
    subtitle = doc.add_paragraph(
        f"Emitido em {date.today().strftime('%d/%m/%Y')} — Aulas registradas: {len(sessions)}{period_txt}"
    )
    for run in subtitle.runs:
        run.font.size = Pt(10)

    # Table: Student | Presencas | Ausencias | % Presenca
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(["Aluno", "Presencas", "Ausencias", "% Presenca"]):
        hdr_cells[i].text = text
    for s in students:
        p = totals[s.id]["present"]
        a = totals[s.id]["absent"]
        pct = (p / (p + a) * 100) if (p + a) else 0
        row_cells = table.add_row().cells
        row_cells[0].text = s.name
        row_cells[1].text = str(p)
        row_cells[2].text = str(a)
        row_cells[3].text = f"{pct:.1f}%"

    doc.add_paragraph("")
    doc.add_paragraph(
        "Este relatorio apresenta a consolidacao de presencas e ausencias por aluno,"
        " com base nos registros disponiveis no periodo."
    )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name=f"Relatorio_{c.name}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@bp.route("/classes/<int:class_id>/students/<int:student_id>/report.docx")
@login_required
@teacher_required
def export_student_docx(class_id: int, student_id: int):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception:
        flash("Dependencia ausente: instale 'python-docx'.", "danger")
        return redirect(url_for("teacher.class_detail", class_id=class_id))

    c = Classroom.query.get_or_404(class_id)
    st = Student.query.get_or_404(student_id)
    if st.class_id != c.id:
        flash("Aluno nao pertence a turma.", "danger")
        return redirect(url_for("teacher.class_detail", class_id=class_id))
    if c.owner_id != current_user.id and current_user.role != "admin":
        flash("Sem permissao.", "danger")
        return redirect(url_for("teacher.dashboard"))

    # Optional stage filter
    stage_id = request.args.get("stage_id", type=int)
    q = Session.query.filter_by(class_id=c.id)
    start = end = None
    if stage_id:
        stg = Stage.query.get_or_404(stage_id)
        start, end = stg.start, stg.end
        q = q.filter(Session.date >= start, Session.date <= end)
    sessions = q.order_by(Session.date).all()
    total_p = total_a = 0
    details = []
    for sess in sessions:
        e = next((x for x in sess.entries if x.student_id == st.id), None)
        present = bool(e.present) if e else False
        total_p += 1 if present else 0
        total_a += 0 if present else 1
        details.append((sess.date, present))
    pct = (total_p / (total_p + total_a) * 100) if (total_p + total_a) else 0

    # Activities and points for the stage (or all)
    acts_q = (
        db.session.query(Activity)
        .join(Stage, Activity.stage_id == Stage.id)
        .filter(Stage.class_id == c.id)
        .order_by(Activity.period_start)
    )
    if stage_id:
        acts_q = acts_q.filter(Activity.stage_id == stage_id)
    activities = acts_q.all()
    # Refresh scores to reflect latest attendance
    for act in activities:
        try:
            _update_scores_for_activity(act, class_id=c.id)
        except Exception:
            pass
    # Build per-activity details and totals
    activities_details = []
    etapa_total_points = 0.0
    for act in activities:
        rows = (
            DailyScore.query.filter_by(activity_id=act.id, student_id=st.id)
            .order_by(DailyScore.date)
            .all()
        )
        day_rows = [(r.date, bool(r.present), float(r.points)) for r in rows]
        act_total = sum(p for _, _, p in day_rows)
        etapa_total_points += act_total
        activities_details.append({
            "title": act.title,
            "period": f"{act.period_start} → {act.period_end}",
            "n": len(day_rows),
            "p_total": float(act.points_total),
            "ppc": float(act.points_per_call),
            "rows": day_rows,
            "act_total": act_total,
        })

    doc = Document()
    section = doc.sections[0]
    section.left_margin, section.right_margin = Inches(0.7), Inches(0.7)
    header = section.header
    h_p = header.paragraphs[0]
    logo = _brand_logo_path()
    if logo:
        run = h_p.add_run()
        try:
            run.add_picture(logo, width=Inches(1.0))
        except Exception:
            pass

    title = doc.add_paragraph()
    r = title.add_run(f"Relatorio Individual — {st.name}")
    r.bold = True
    r.font.size = Pt(16)
    period_txt = f"  •  Período: {start} → {end}" if start and end else ""
    subtitle = doc.add_paragraph(
        f"Turma: {c.name}  •  Presencas: {total_p}  •  Ausencias: {total_a}  •  % Presenca: {pct:.1f}%{period_txt}"
    )
    for run in subtitle.runs:
        run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Data"
    hdr[1].text = "Registro"
    for d, present in details:
        row = table.add_row().cells
        row[0].text = d
        row[1].text = "Presente" if present else "Ausente"

    doc.add_paragraph("")
    # Add activities section with points
    doc.add_paragraph("Atividades da Etapa")
    for info in activities_details:
        p = doc.add_paragraph()
        p.add_run(f"• {info['title']} — Período: {info['period']}  | N: {info['n']}  | ")
        p.add_run(f"P_total: {info['p_total']:.2f}  | P/ chamada: {info['ppc']:.2f}")
        t = doc.add_table(rows=1, cols=3)
        h = t.rows[0].cells
        h[0].text = "Dia"; h[1].text = "Presença"; h[2].text = "Pontos"
        for d, pres, pts in info["rows"]:
            rw = t.add_row().cells
            rw[0].text = d
            rw[1].text = "P" if pres else "A"
            rw[2].text = f"{pts:.2f}"
        doc.add_paragraph(f"Total na atividade: {info['act_total']:.2f} ponto(s)")
        doc.add_paragraph("")

    doc.add_paragraph(f"Total da etapa (pontos): {etapa_total_points:.2f}")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    safe_name = st.name.replace(" ", "_")
    return send_file(
        output,
        as_attachment=True,
        download_name=f"Relatorio_{safe_name}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
